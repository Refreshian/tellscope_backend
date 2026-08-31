import asyncio
import json
import os
import re
import sys
import time
from typing import List, Dict, Any

import pandas as pd
import aiohttp
import matplotlib.pyplot as plt
from io import BytesIO

from docx import Document
from docx.shared import Inches

from mlops.gateway import GatewayChatClient, achat
from mlops.lock import external_cfg, generate_cfg, prompt_id as lock_prompt_id
from mlops.prompts import render_prompt

# ==========================================
# КОНФИГУРАЦИЯ (модели из serving lock, ключи только из .env)
# ==========================================
EXTERNAL_MODEL = (external_cfg("smart_agent_planner").get("model") or "claude-sonnet-4.5")
LOCAL_MODEL_NAME = (generate_cfg().get("model") or "Qwen/Qwen3-32B-FP8")
BATCH_SIZE = 32
MAX_CONCURRENCY = 64
CONNECT_TIMEOUT = 15
TOTAL_TIMEOUT = 180
MAX_RETRIES = 2

# ==========================================
# 1. МОДУЛЬ ВНЕШНЕЙ LLM (Оркестратор) - ЗНАЧИТЕЛЬНЫЕ УЛУЧШЕНИЯ
# ==========================================

class ExternalLLMbrain:
    def __init__(self):
        self.client = GatewayChatClient(provider="aitunnel", profile="smart_agent_planner")

    def _is_english_requested(self, text: str) -> bool:
        lowered = text.lower()
        return "english" in lowered or "англий" in lowered

    def plan_task(self, user_query: str, data_sample: str, available_columns: List[str] = None) -> Dict[str, Any]:
        """
        Строит JSON-план анализа данных по запросу пользователя.
        План должен описывать:
        - какие режимы анализа нужны (темы, тональность, демография);
        - какие графики и по каким полям нужно построить;
        - структуру итогового Word-отчета.
        """
        columns_info = f"\n\nДоступные колонки в данных: {', '.join(available_columns)}" if available_columns else ""
        needs_english = self._is_english_requested(user_query)
        target_language = "английском" if needs_english else "русском"

        try:
            system_prompt = render_prompt(
                lock_prompt_id("smart_agent_plan", "smart_agent_plan_v1"),
                columns_info=columns_info,
                target_language=target_language,
            )
        except Exception as e:
            print(f"ОШИБКА загрузки промпта плана: {e}. Используется план по умолчанию.")
            return self._get_default_plan(user_query, target_language)
        try:
            response = self.client.chat.completions.create(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Запрос: {user_query}\n\nСэмпл данных:\n{data_sample}"}
                ],
                model=EXTERNAL_MODEL,
                temperature=0.1,
                response_format={"type": "json_object"}
            )
            content = response.choices[0].message.content
            print(f"DEBUG: Получен план от API:\n{content}\n")
            plan = json.loads(content)
            from mlops.eval_plan import validate_plan
            errors = validate_plan(plan)
            if errors:
                raise ValueError("; ".join(errors))
            return plan
        except Exception as e:
            print(f"ОШИБКА при создании плана: {e}. Используется план по умолчанию.")
            return self._get_default_plan(user_query, target_language)

    def generate_thematic_intro(self, topic_name: str, topic_examples: List[str], user_query: str) -> str:
        examples_str = "\n".join([f"- {ex[:250]}..." for ex in topic_examples])
        try:
            prompt = render_prompt(
                lock_prompt_id("smart_agent_intro", "smart_agent_intro_v1"),
                topic_name=topic_name,
                user_query=user_query,
                examples_str=examples_str,
            )
        except Exception:
            prompt = (
                f"Напиши короткое введение (3-4 предложения) для тематики \"{topic_name}\" "
                f"в контексте \"{user_query}\". Примеры:\n{examples_str}"
            )
        try:
            response = self.client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model=EXTERNAL_MODEL,
                temperature=0.4
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f"ОШИБКА при генерации введения для темы '{topic_name}': {e}")
            return f"Не удалось автоматически сгенерировать описание для тематики '{topic_name}'. Эта тема выделена на основе анализа сообщений."

    def generate_narrative(self, stats_data: str, user_query: str) -> str:
        try:
            prompt = render_prompt(
                lock_prompt_id("smart_agent_narrative", "smart_agent_narrative_v1"),
                user_query=user_query,
                stats_data=stats_data,
            )
        except Exception:
            prompt = f"Напиши итоговые выводы по запросу \"{user_query}\". Статистика:\n{stats_data}"
        try:
            response = self.client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model=EXTERNAL_MODEL,
                temperature=0.5
            )
            content = response.choices[0].message.content
            return content.strip()
        except Exception as e:
            print(f"ОШИБКА при генерации выводов: {e}")
            return f"Автоматическая генерация выводов недоступна. Статистика:\n{stats_data}"

    def _get_default_plan(self, user_query: str, target_language: str = "русском") -> Dict[str, Any]:
        return {
            "filters": {},
            "analysis_needed": True,
            "local_llm_system_prompt": f"Ты — классификатор текста. Определи основную тему сообщения. Ответ - ТОЛЬКО название темы (1-3 слова) на {target_language}. Без рассуждений.",
            "local_llm_user_question": "Определи тему этого сообщения.",
            "report_title": f"Анализ по запросу: {user_query[:50]}",
            "report_structure": [
                {"type": "title"},
                {"type": "section", "title": "1. Введение", "content_source": "user_query"},
                {"type": "overall_stats", "title": "Общая статистика"},
                {
                    "type": "thematic_breakdown",
                    "title": "2. Анализ ключевых тематик",
                    "breakdown_by_column": "llm_result_clean",
                    "max_themes": 5, "num_examples": 5,
                    "thematic_section_structure": ["introduction", "demographics_chart", "examples"]
                },
                {"type": "section", "title": "3. Итоговые выводы", "content_source": "final_conclusions"}
            ]
        }
# ... Классы LocalLLMWorker и ReportBuilder остаются без изменений ...
class LocalLLMWorker:
    def __init__(self): self.session = None
    async def _generate_single(self, text: str, system_prompt: str, user_question: str) -> str:
        if not text: return "Пустой текст"
        cleaned_text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)[:4000]
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"{user_question}\n\nТекст:\n{cleaned_text}"},
        ]
        for attempt in range(MAX_RETRIES + 1):
            try:
                result = await achat(
                    provider="vllm",
                    messages=messages,
                    model=LOCAL_MODEL_NAME,
                    temperature=0.1,
                    max_tokens=200,
                    timeout=TOTAL_TIMEOUT,
                )
                answer = (result.content or "").strip()
                return re.sub(r'(?i)\bthink\b[:：]?', '', answer).split("Ответ:")[-1].strip()
            except Exception:
                await asyncio.sleep(0.5 * (2 ** attempt))
        return "Ошибка анализа"
    async def process_batch(self, texts: List[str], system_prompt: str, user_question: str) -> List[str]:
        semaphore = asyncio.Semaphore(MAX_CONCURRENCY)
        async def worker(txt):
            async with semaphore: return await self._generate_single(txt, system_prompt, user_question)
        tasks = [asyncio.create_task(worker(t)) for t in texts]; total = len(tasks)
        for i, f in enumerate(asyncio.as_completed(tasks)):
            if i % 10 == 0:
                await self.progress_callback(f"Анализ текстов: {i}/{total}...")
            await f
        return await asyncio.gather(*tasks)

class ReportBuilder:
    def __init__(self, filename="Analytics_Report.docx"): self.doc = Document(); self.filename = filename
    def add_title(self, text): self.doc.add_heading(text, 0)
    def add_heading(self, text, level=1): self.doc.add_heading(text, level=level)
    def add_paragraph(self, text): self.doc.add_paragraph(text)
    def add_section(self, title, content, level=1): self.add_heading(title, level); self.add_paragraph(content)
    def add_list_stats(self, title, stats_dict: Dict, level=2):
        self.add_heading(title, level)
        for key, value in stats_dict.items(): self.doc.add_paragraph(f"{key}: {value}", style='List Bullet')
    def add_examples(self, examples: List[Dict], title="Примеры сообщений", level=3):
        self.add_heading(title, level)
        for ex in examples:
            p = self.doc.add_paragraph(); p.add_run(f"Автор: {ex.get('authorObject', {}).get('fullname', 'N/A')}, Источник: {ex.get('hub', 'N/A')}\n").bold = True
            p.add_run(ex.get('text', '')[:300] + "...\n"); p.add_run(f"Ссылка: {ex.get('url', '#')}\n").italic = True
    def add_chart(self, title: str, data: pd.Series, chart_type: str = "bar", level=2):
        # ИСПРАВЛЕНИЕ 1: Проверяем, есть ли вообще данные
        if data.empty or data.sum() == 0:
            self.add_heading(title, level)
            self.add_paragraph(f"На графике '{title}' нет данных для отображения.")
            return

        self.add_heading(title, level)
        plt.figure(figsize=(6.5, 4))

        # Отрисовка графика (без изменений)
        if chart_type == "bar":
            data.plot(kind='bar', color='skyblue')
        elif chart_type == 'pie':
            # Для pie чартов лучше отфильтровать очень маленькие значения
            data_for_pie = data[data > 0]
            data_for_pie.plot(kind='pie', autopct='%1.1f%%', startangle=90)
            plt.ylabel('')
        else:
            data.plot(kind='bar')

        plt.title(title)
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        plt.close()
        img_stream.seek(0)
        self.doc.add_picture(img_stream, width=Inches(6))

        # ИСПРАВЛЕНИЕ 2: Полностью переработанная логика объяснения
        explanation = f"На графике '{title}' представлено распределение авторов."
        
        # Проверяем, есть ли реальные данные для анализа
        if not data.empty and data.max() > 0:
            max_value = data.max()
            max_category = data.idxmax()
            # Заменяем пустую категорию на более понятное описание
            if str(max_category).strip() == '':
                max_category = 'Не указан'

            if chart_type == 'pie':
                total = data.sum()
                percentage = (max_value / total) if total > 0 else 0
                explanation += f" Наибольшую долю ({percentage:.1%}) составляет категория '{max_category}' ({int(max_value)} авторов)."
            else:
                explanation += f" Самой многочисленной категорией является '{max_category}' с {int(max_value)} авторами."
        else:
            explanation += " Нет данных для определения доминирующей категории."
            
        self.add_paragraph(explanation)

    def save(self): self.doc.save(self.filename); return self.filename

# ==========================================
# 4. ГЛАВНЫЙ КЛАСС АГЕНТА - ВСЯ ЛОГИКА ЗДЕСЬ
# ==========================================

class SocialMediaAgent:
    def __init__(self, progress_callback=None):
        self.brain = ExternalLLMbrain()
        self.worker = LocalLLMWorker()
        self.worker.progress_callback = progress_callback # Передаем callback в воркер
        self.data_df = None
        # Кэш плана на время выполнения задачи
        self._plan: Dict[str, Any] = {}
        self.progress_callback = progress_callback

    async def _log_progress(self, message: str):
        if self.progress_callback: await self.progress_callback(message)
        print(message)

    def load_data(self, json_path):
        with open(json_path, 'r', encoding='utf-8') as f: data = json.load(f)
        self.data_df = pd.DataFrame(data['items'] if isinstance(data, dict) and 'items' in data else data)
        # Очистка и обогащение данных
        if not self.data_df.empty:
            self.data_df['sex'] = self.data_df['authorObject'].apply(lambda x: x.get('sex') if isinstance(x, dict) else None)
            self.data_df['age'] = self.data_df['authorObject'].apply(lambda x: x.get('age') if isinstance(x, dict) else None)
            self.data_df['age'] = pd.to_numeric(self.data_df['age'], errors='coerce')
            self.data_df['age'] = self.data_df['age'].where((self.data_df['age'] >= 0) & (self.data_df['age'] <= 120))
            bins = [0, 18, 25, 35, 45, 60, 120]; labels = ['<18', '18-24', '25-34', '35-44', '45-59', '60+']
            self.data_df['age_group'] = pd.cut(self.data_df['age'], bins=bins, labels=labels, right=False)

    def filter_data(self, filters: Dict):
        df = self.data_df.copy()
        # Тут может быть сложная логика фильтрации
        self.data_df = df

    def _get_analysis_modes(self, plan: Dict[str, Any]) -> Dict[str, Any]:
        """
        Безопасно читает блок analysis_modes с разумными значениями по умолчанию
        для обратной совместимости со старыми планами.
        """
        default_modes = {
            "topics": plan.get("analysis_needed", True),
            "sentiment": False,
            "demographics": ["sex", "age_group", "hub"],
        }
        modes = plan.get("analysis_modes") or {}
        # не ломаемся, если пришёл странный тип
        if not isinstance(modes, dict):
            return default_modes
        return {
            "topics": bool(modes.get("topics", default_modes["topics"])),
            "sentiment": bool(modes.get("sentiment", default_modes["sentiment"])),
            "demographics": modes.get("demographics", default_modes["demographics"]),
        }

    async def _run_sentiment_analysis_if_needed(self, analysis_modes: Dict[str, Any]):
        """
        Запускает дополнительный проход локальной LLM для тональности,
        если это указано в analysis_modes.
        """
        if not analysis_modes.get("sentiment"):
            return

        if self.data_df is None or self.data_df.empty:
            return

        await self._log_progress("Этап 1b: Анализ тональности сообщений...")

        # Промпты можно позже сделать настраиваемыми через план, пока — дефолтные.
        system_prompt = (
            "Ты — анализатор тональности текста. "
            "Определи общую тональность сообщения: Позитив, Негатив или Нейтрально. "
            "Ответь ТОЛЬКО одним словом из этого списка."
        )
        user_question = "Определи тональность этого сообщения."

        texts = self.data_df["text"].fillna("").tolist()
        sentiments = await self.worker.process_batch(texts, system_prompt, user_question)

        def _normalize_sentiment(raw: str) -> str:
            if not raw:
                return "Не определено"
            t = raw.strip().lower()
            if "позит" in t or "positive" in t:
                return "Позитив"
            if "негат" in t or "negative" in t:
                return "Негатив"
            if "нейтр" in t or "neutral" in t:
                return "Нейтрально"
            return "Не определено"

        self.data_df["sentiment"] = [ _normalize_sentiment(x) for x in sentiments ]

    def _get_chart_defs(self, plan: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
        """
        Возвращает словарь описаний графиков по их id.
        Если в плане нет поля charts, добавляет разумный дефолтный набор.
        """
        charts_list = plan.get("charts")
        if not charts_list or not isinstance(charts_list, list):
            charts_list = [
                {
                    "id": "sex_distribution",
                    "title": "Распределение авторов по полу",
                    "type": "pie",
                    "group_by": "sex",
                    "filter": {},
                },
                {
                    "id": "age_distribution",
                    "title": "Распределение авторов по возрастным группам",
                    "type": "bar",
                    "group_by": "age_group",
                    "filter": {},
                },
                {
                    "id": "hub_distribution",
                    "title": "Распределение сообщений по площадкам",
                    "type": "bar",
                    "group_by": "hub",
                    "filter": {},
                },
            ]
        charts_map = {}
        for ch in charts_list:
            if isinstance(ch, dict) and "id" in ch:
                charts_map[ch["id"]] = ch
        return charts_map

    def _build_series_for_chart(self, chart_def: Dict[str, Any], df: pd.DataFrame) -> pd.Series:
        """
        Универсальный билдер pd.Series для графика по описанию chart_def.
        Умеет учитывать простейший фильтр (равенства по колонкам).
        """
        if df is None or df.empty:
            return pd.Series(dtype=float)

        work_df = df
        filter_spec = chart_def.get("filter") or {}
        if isinstance(filter_spec, dict) and filter_spec:
            for col, value in filter_spec.items():
                if col in work_df.columns:
                    work_df = work_df[work_df[col] == value]

        group_by = chart_def.get("group_by")
        if not group_by or group_by not in work_df.columns:
            return pd.Series(dtype=float)

        series = work_df[group_by].astype(str).replace(
            ["nan", "None", ""], "Не указан"
        ).value_counts()
        return series

    async def run_task(self, user_query: str, input_file: str, report_save_path: str):
        # 1. ЗАГРУЗКА И ПЛАНИРОВАНИЕ
        await self._log_progress("Загрузка и подготовка данных...")
        self.load_data(input_file)
        sample = self.data_df.head(1).to_json(orient='records', force_ascii=False) if not self.data_df.empty else "{}"
        columns = self.data_df.columns.tolist() if not self.data_df.empty else []
        
        await self._log_progress("Планирование анализа (мозг агента)...")
        plan = await asyncio.to_thread(self.brain.plan_task, user_query, sample, columns)
        # сохраняем план локально, чтобы переиспользовать
        self._plan = plan or {}
        await self._log_progress(f"План утвержден: {json.dumps(self._plan, ensure_ascii=False, indent=2)}")

        analysis_modes = self._get_analysis_modes(self._plan)
        charts_map = self._get_chart_defs(self._plan)

        # 2. ФИЛЬТРАЦИЯ
        await self._log_progress("Применение фильтров...")
        self.filter_data(self._plan.get('filters', {}))
        if self.data_df.empty:
            raise Exception("Нет данных после фильтрации. Отчет не может быть создан.")

        # 3. АНАЛИЗ ТЕМ (КЛАССИФИКАЦИЯ)
        if analysis_modes.get("topics") and self._plan.get("analysis_needed"):
            await self._log_progress("Этап 1: Классификация сообщений по тематикам...")
            texts = self.data_df["text"].fillna("").tolist()
            labels = await self.worker.process_batch(
                texts,
                self._plan["local_llm_system_prompt"],
                self._plan["local_llm_user_question"],
            )
            self.data_df["llm_result_clean"] = [
                re.sub(r"[^\w\s-]", "", str(x)).strip().capitalize() for x in labels
            ]
            self.data_df = self.data_df[self.data_df["llm_result_clean"].str.len() > 1]
            self.data_df = self.data_df[
                ~self.data_df["llm_result_clean"].str.contains(
                    "Ошибка анализа", case=False
                )
            ]

        # 3b. АНАЛИЗ ТОНАЛЬНОСТИ (ОПЦИОНАЛЬНО)
        await self._run_sentiment_analysis_if_needed(analysis_modes)
        
        # 4. СБОРКА ОТЧЕТА ПО ПЛАНУ
        await self._log_progress("Этап 2: Сборка итогового отчета...")
        report = ReportBuilder(filename=report_save_path)
        
        report_context = {
            "user_query": user_query,
            "report_title": self._plan.get("report_title", "Аналитический отчет"),
        }

        for section_plan in self._plan.get("report_structure", []):
            section_type = section_plan.get("type")
            
            if section_type == "title": report.add_title(report_context["report_title"])
            
            elif section_type == "section" and section_plan.get("content_source") == "user_query":
                report.add_section(section_plan.get("title", "Введение"), f'Отчет подготовлен по запросу: "{user_query}"')
                
            elif section_type == "overall_stats":
                await self._log_progress("Подсчет общей статистики...")
                total_messages = len(self.data_df); unique_authors = self.data_df['authorObject'].apply(lambda x: x.get('id') if isinstance(x, dict) else None).nunique()
                top_hubs = self.data_df['hub'].value_counts().head(5)
                report.add_heading(section_plan.get("title", "Общая статистика"), level=1)
                report.add_paragraph(f"Всего проанализировано сообщений: {total_messages}\nУникальных авторов: {unique_authors}")
                report.add_chart("Топ-5 площадок по сообщениям", top_hubs, chart_type='bar', level=2)
                
            elif section_type == "thematic_breakdown" and 'llm_result_clean' in self.data_df.columns:
                await self._log_progress("Начинаю детальный анализ по тематикам...")
                report.add_heading(section_plan.get("title", "Детальный анализ тематик"), level=1)
                
                breakdown_col = section_plan.get("breakdown_by_column", "llm_result_clean")
                max_themes = section_plan.get("max_themes", 7)
                topics_to_analyze = self.data_df[breakdown_col].value_counts().head(max_themes)
                
                for i, (topic, count) in enumerate(topics_to_analyze.items()):
                    await self._log_progress(f"Анализ темы {i+1}/{len(topics_to_analyze)}: '{topic}'")
                    topic_df = self.data_df[self.data_df[breakdown_col] == topic]
                    report.add_heading(f"{i+1}. Тематика: {topic} ({count} сообщ.)", level=2)
                    
                    for sub_task in section_plan.get("thematic_section_structure", []):
                        if sub_task == "introduction":
                            examples = topic_df['text'].head(5).tolist()
                            intro_text = await asyncio.to_thread(self.brain.generate_thematic_intro, topic, examples, user_query)
                            report.add_paragraph(intro_text)
                            
                        if sub_task == "demographics_chart":
                            sex_counts = topic_df['sex'].astype(str).replace(['nan', 'None', ''], 'Не указан').value_counts()
                            # Убираем "Не указан" из pie-чарта, если он неинформативен, или оставляем, если это основная масса
                            if 'Не указан' in sex_counts and len(sex_counts) > 1:
                                # Можно решить, показывать ли "Не указан"
                                pass # Пока оставляем для полноты картины
                            sex_dist = sex_counts

                            age_dist = topic_df['age_group'].dropna().value_counts().sort_index()
                            
                            report.add_chart(f"Распределение по полу (тема: {topic})", sex_dist, chart_type='pie', level=3)
                            report.add_chart(f"Распределение по возрасту (тема: {topic})", age_dist, chart_type='bar', level=3)

                        if sub_task == "examples":
                            num_examples = section_plan.get("num_examples", 7)
                            examples_data = topic_df.head(num_examples).to_dict('records')
                            report.add_examples(examples_data, title="Примеры сообщений по теме", level=3)

            elif section_type == "sentiment_overview" and "sentiment" in self.data_df.columns:
                await self._log_progress("Формирование раздела с анализом тональности...")
                title = section_plan.get("title", "Анализ тональности")
                report.add_heading(title, level=1)

                sentiment_col = section_plan.get("sentiment_column", "sentiment")
                if sentiment_col in self.data_df.columns:
                    sent_dist = (
                        self.data_df[sentiment_col]
                        .astype(str)
                        .replace(["nan", "None", ""], "Не определено")
                        .value_counts()
                    )
                    report.add_chart(
                        "Распределение сообщений по тональности",
                        sent_dist,
                        chart_type="pie",
                        level=2,
                    )

            elif section_type == "custom_chart_section":
                await self._log_progress("Формирование пользовательского раздела с графиками...")
                title = section_plan.get("title", "Дополнительные графики")
                report.add_heading(title, level=1)

                chart_ids = section_plan.get("charts") or []
                for ch_id in chart_ids:
                    ch_def = charts_map.get(ch_id)
                    if not ch_def:
                        continue
                    series = self._build_series_for_chart(ch_def, self.data_df)
                    chart_title = ch_def.get("title", ch_id)
                    chart_type = ch_def.get("type", "bar")
                    report.add_chart(chart_title, series, chart_type=chart_type, level=2)

        # 5. ГЕНЕРАЦИЯ ИТОГОВЫХ ВЫВОДОВ
        final_conclusion_section = next(
            (s for s in self._plan.get("report_structure", []) if s.get("content_source") == "final_conclusions"),
            None,
        )
        if final_conclusion_section:
            await self._log_progress("Этап 3: Генерация итоговых выводов...")
            # Расширенный набор статистик для LLM
            stats_summary = {
                "total_messages": len(self.data_df),
                "unique_authors": self.data_df["authorObject"]
                .apply(lambda x: x.get("id") if isinstance(x, dict) else None)
                .nunique(),
                "analysis_top_results": self.data_df.get(
                    "llm_result_clean", pd.Series(dtype=str)
                )
                .value_counts()
                .head(5)
                .to_dict(),
            }
            if "sentiment" in self.data_df.columns:
                stats_summary["sentiment_distribution"] = (
                    self.data_df["sentiment"]
                    .astype(str)
                    .replace(["nan", "None", ""], "Не определено")
                    .value_counts()
                    .to_dict()
                )
            if "hub" in self.data_df.columns:
                stats_summary["hub_distribution"] = (
                    self.data_df["hub"].astype(str).value_counts().head(10).to_dict()
                )
            stats_str = json.dumps(stats_summary, ensure_ascii=False)
            conclusions = await asyncio.to_thread(self.brain.generate_narrative, stats_str, user_query)
            report.add_section(final_conclusion_section.get("title", "Итоговые выводы"), conclusions, level=1)

        # 6. СОХРАНЕНИЕ
        saved_name = report.save()
        await self._log_progress(f"✅ Готово! Отчет сохранен: {saved_name}")
        return saved_name

async def main():
    if not os.path.exists("data.json"): print("Файл data.json не найден."); return
    user_request = 'Проанализируй все сообщения темы "Платон - система взимания оплаты проезда", создай итоговый отчет Word на основе обсуждений в соцмедиа: какие тематики есть и сделай отчет по каждой тематике: 1. Во введении каждой тематики расскажи про что эта тематика и как ее обсуждают, 2. Нарисуй и объясни график распределения авторов по тематике: по полу, возрасту. 3. Приведи (5-7) примеров сообщений по каждой тематике. 4. Сделай выводы по отчету обсуждений'
    agent = SocialMediaAgent(progress_callback=lambda msg: print(f"[PROGRESS] {msg}"))
    try: await agent.run_task(user_request, "data.json", f"Detailed_Report_{int(time.time())}.docx")
    except Exception as e: print(f"Критическая ошибка: {e}"); import traceback; traceback.print_exc()

if __name__ == "__main__":
    if sys.platform == 'win32': asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
    asyncio.run(main())